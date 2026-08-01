"""设计算法文档导入：Markdown / Word(.docx) / Excel(.xlsx) → algorithms/<module>.yaml。"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List, Optional

import yaml


STEP_HEADERS = {"步骤", "操作", "算法", "说明", "描述", "内容", "label", "desc", "step"}
KIND_HEADERS = {"类型", "类别", "kind", "type"}
INPUT_HEADERS = {"输入", "inputs", "input"}
OUTPUT_HEADERS = {"输出", "outputs", "output"}
NOTE_HEADERS = {"备注", "注意", "说明", "note"}


def guess_kind(text: str) -> str:
    t = str(text or "")
    if re.search(r"输入|读取|加载|查询|解析|获取|拉取|read|load|parse|get|fetch|query|search", t, re.I):
        return "read"
    if re.search(r"校验|判断|检查|权限|过滤|如果|校验|check|verify|validate|if|filter|authoriz|permission", t, re.I):
        return "control"
    if re.search(r"输出|写入|渲染|生成|保存|导出|写文件|output|write|render|export|save", t, re.I):
        return "output"
    if re.search(r"调用外部|外部服务|http|api|rpc|feign", t, re.I):
        return "external"
    return "transform"


def _split_list(text: str) -> List[str]:
    return [s.strip() for s in re.split(r"[、,，;；/]", text) if s.strip()]


def _parse_details(label: str) -> dict:
    """从步骤文本提取 输入/输出/备注。"""
    inputs: List[str] = []
    outputs: List[str] = []
    note = ""
    m = re.search(r"[（(]([^）)]*)[）)]", label)
    if m:
        for p in re.split(r"[，,]", m.group(1)):
            p = p.strip()
            if p.startswith("输入"):
                inputs = _split_list(p.lstrip("输入:： "))
            elif p.startswith("输出"):
                outputs = _split_list(p.lstrip("输出:： "))
            elif p.startswith(("备注", "注意")):
                note = p.lstrip("备注注意:： ").strip()
    return {"inputs": inputs, "outputs": outputs, "note": note}


def _norm_step(idx: int, label: str, kind: str = "", inputs: Optional[List[str]] = None,
               outputs: Optional[List[str]] = None, note: str = "") -> dict:
    label = re.sub(r"[（(]\s*(?:输入|输出|备注)[^）)]*[）)]", "", str(label)).strip()
    label = re.sub(r"^\s*\d+[.、)）]\s*", "", label).strip("。 .")
    return {
        "id": f"step-{idx + 1}",
        "kind": kind or guess_kind(label),
        "label": label or f"步骤 {idx + 1}",
        "inputs": [str(i) for i in (inputs or [])],
        "outputs": [str(o) for o in (outputs or [])],
        "note": note or "",
    }


def import_markdown(path: str | Path) -> List[dict]:
    lines = Path(path).read_text(encoding="utf-8").splitlines()
    steps: List[dict] = []
    header: Optional[List[str]] = None
    current: Optional[dict] = None
    pending_label = ""

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        # 表格
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if header is None:
                header = cells
                continue
            sep = cells[0].strip().replace("-", "").replace(":", "")
            if sep == "" and cells[0].strip():
                continue
            if all(not c for c in cells):
                continue
            steps.append(_table_row_to_step(len(steps), cells, header))
            current = None
            continue
        header = None

        m = re.match(r"^\s*(?:\d+[.、)）]|[-*])\s+(.+)$", line)
        heading = re.match(r"^#{2,4}\s+(.+)$", line)
        if heading:
            if current and pending_label:
                current["label"] = pending_label
            current = None
            pending_label = heading.group(1).strip()
            continue
        if m:
            if current and pending_label:
                current["label"] = pending_label
            text = m.group(1).strip()
            details = _parse_details(text)
            step = _norm_step(len(steps), text, inputs=details["inputs"],
                              outputs=details["outputs"], note=details["note"])
            steps.append(step)
            current = step
            pending_label = ""
            continue
        # 细节行（输入:/输出:/备注:）附加到当前步骤
        if current:
            m_in = re.match(r"^\s*输入[:：]\s*(.+)$", line)
            m_out = re.match(r"^\s*输出[:：]\s*(.+)$", line)
            m_note = re.match(r"^\s*(?:备注|注意)[:：]\s*(.+)$", line)
            if m_in:
                current["inputs"] = _split_list(m_in.group(1))
            elif m_out:
                current["outputs"] = _split_list(m_out.group(1))
            elif m_note:
                current["note"] = m_note.group(1).strip()
    if current and pending_label:
        current["label"] = pending_label
    return steps


def _table_row_to_step(idx: int, cells: List[str], header: List[str]) -> dict:
    def col(*names) -> Optional[str]:
        for i, h in enumerate(header):
            if any(n.lower() == h.strip().lower() for n in names):
                return cells[i] if i < len(cells) else ""
        return None

    label = col("步骤", "操作", "算法", "说明", "描述", "内容", "label", "desc") or ""
    kind = col("类型", "类别", "kind", "type") or ""
    inputs = col("输入", "inputs") or ""
    outputs = col("输出", "outputs") or ""
    note = col("备注", "注意", "说明", "note") or ""
    return _norm_step(idx, label, kind=kind,
                      inputs=_split_list(inputs) if inputs else None,
                      outputs=_split_list(outputs) if outputs else None,
                      note=note)


def import_docx(path: str | Path) -> List[dict]:
    from docx import Document

    doc = Document(str(path))
    steps: List[dict] = []
    current: Optional[dict] = None

    def flush():
        nonlocal current
        current = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        is_list = "list" in style or bool(re.match(r"^\s*(?:\d+[.、)）]|[-*•])\s+", text))
        if is_list:
            if current:
                current = None
            details = _parse_details(text)
            step = _norm_step(len(steps), text, inputs=details["inputs"],
                              outputs=details["outputs"], note=details["note"])
            steps.append(step)
            current = step
            continue
        if current:
            m_in = re.match(r"^\s*输入[:：]\s*(.+)$", text)
            m_out = re.match(r"^\s*输出[:：]\s*(.+)$", text)
            m_note = re.match(r"^\s*(?:备注|注意)[:：]\s*(.+)$", text)
            if m_in:
                current["inputs"] = _split_list(m_in.group(1))
            elif m_out:
                current["outputs"] = _split_list(m_out.group(1))
            elif m_note:
                current["note"] = m_note.group(1).strip()

    for table in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        if not rows:
            continue
        header = rows[0]
        if not any(any(n.lower() == h.lower() for n in
                       STEP_HEADERS | KIND_HEADERS | INPUT_HEADERS | OUTPUT_HEADERS)
                   for h in header):
            continue
        for row in rows[1:]:
            steps.append(_table_row_to_step(len(steps), row, header))
    return steps


def import_xlsx(path: str | Path) -> List[dict]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    ws = wb.worksheets[0]
    rows = [list(r) for r in ws.iter_rows(values_only=True)]
    rows = [[str(c) if c is not None else "" for c in row] for row in rows]
    wb.close()
    if not rows:
        return []

    header_idx = None
    for i, row in enumerate(rows[:5]):
        if any(any(n.lower() == h.strip().lower() for n in
                   STEP_HEADERS | KIND_HEADERS | INPUT_HEADERS | OUTPUT_HEADERS | NOTE_HEADERS)
               for h in row):
            header_idx = i
            break
    if header_idx is None:
        # 无表头：整行视为"说明"
        return [_norm_step(i, row[0] or " ".join(x for x in row if x))
                for i, row in enumerate(rows) if any(row)]

    header = rows[header_idx]
    steps = []
    for row in rows[header_idx + 1:]:
        if not any(row):
            continue
        steps.append(_table_row_to_step(len(steps), row, header))
    return steps


def import_design(path: str | Path) -> List[dict]:
    suffix = Path(path).suffix.lower()
    if suffix == ".md" or suffix == ".markdown":
        return import_markdown(path)
    if suffix == ".docx":
        return import_docx(path)
    if suffix == ".xlsx":
        return import_xlsx(path)
    raise ValueError(f"不支持的设计文档格式: {suffix}（支持 .md/.docx/.xlsx）")


def write_spec(module: str, steps: List[dict], out_path: str | Path,
               description: str = "") -> Path:
    spec = {
        "module": module,
        "description": description or f"{module} 设计算法（导入自文档）",
        "steps": steps,
    }
    p = Path(out_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(yaml.safe_dump(spec, allow_unicode=True, sort_keys=False),
                 encoding="utf-8")
    return p
