"""端到端冒烟测试：在 fixture 上跑通 index → discover → resolve → export → validate。"""

from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from rs.config import load_config
from rs.align import align_algorithms, load_design_spec
from rs.algorithm import reconstruct_algorithm
from rs.pipeline import run_module
from rs.store import load_graph, out_dir


FIXTURE = ROOT / "tests" / "fixtures" / "java-spring-demo"


def reset(project_dir: Path) -> None:
    for d in ("out", ".report-spliter"):
        shutil.rmtree(project_dir / d, ignore_errors=True)


def main() -> int:
    reset(FIXTURE)
    cfg = load_config(FIXTURE)

    # 模块划分结果（期望的文件集合）
    expected = {
        "report-export": {
            "controller/ReportExportController.java",
            "service/ReportExportService.java",
            "service/ReportCalculationService.java",
            "service/PdfRenderer.java",
            "repo/ReportRepository.java",
            "model/Report.java",
            "support/DateUtils.java",
            "support/Metrics.java",
        },
        "user-mgmt": {
            "controller/UserController.java",
            "service/UserService.java",
            "repo/UserRepository.java",
            "model/User.java",
            "support/DateUtils.java",
        },
    }
    prefix = "src/main/java/com/acme/report/"

    failed = False
    for module, want in expected.items():
        summary = run_module(cfg, module, quiet=True)
        manifest = json.loads((out_dir(FIXTURE, module) / "manifest.json").read_text(encoding="utf-8"))
        got = {
            f["path"].removeprefix(prefix)
            for f in manifest["files"]
            if f["kind"] == "source"
        }
        missing = want - got
        extra = got - want
        status = "PASS" if not missing and not extra else "FAIL"
        if status == "FAIL":
            failed = True
        print(f"[{status}] {module}: {len(got)} 文件")
        if missing:
            print(f"    缺失: {sorted(missing)}")
        if extra:
            print(f"    多余: {sorted(extra)}")
        print(f"    审计: {summary['audit_stats']}")

    # 交叉污染检查：report-export 不能含 User，user-mgmt 不能含 Report
    for module, forbidden in (("report-export", {"User"}), ("user-mgmt", {"Report"})):
        graph = load_graph(FIXTURE, module)
        leak = [s for s in graph.symbols if any(f in s for f in forbidden)
                and "ReportApplication" not in s]
        if leak:
            failed = True
            print(f"[FAIL] {module} 泄漏符号: {leak[:10]}")
        else:
            print(f"[PASS] {module} 无交叉泄漏")

    # 共享代码检查：DateUtils 在两个模块中都被纳入（允许重复）
    for module in ("report-export", "user-mgmt"):
        graph = load_graph(FIXTURE, module)
        if not any("DateUtils" in s for s in graph.symbols):
            failed = True
            print(f"[FAIL] {module} 缺少共享代码 DateUtils")
        else:
            print(f"[PASS] {module} 包含共享代码 DateUtils")

    print("\n" + ("ALL PASS" if not failed else "HAS FAILURES"))
    return 0 if not failed else 1


def test_algorithm_and_alignment() -> int:
    reset(FIXTURE)
    cfg = load_config(FIXTURE)
    summary = run_module(cfg, "report-export", quiet=True)
    from rs.store import load_index

    index = load_index(FIXTURE)
    graph = load_graph(FIXTURE, "report-export")
    alg = reconstruct_algorithm(index, graph, FIXTURE)

    failed = False
    steps = [s for e in alg["entries"] for s in e["steps"]]
    anchored = all(s.get("file") and s.get("line") for s in steps)
    if len(alg["entries"]) != 2:
        failed = True
        print(f"[FAIL] 算法重建入口数 {len(alg['entries'])} != 2")
    elif not anchored:
        failed = True
        print("[FAIL] 存在缺少代码锚点的步骤")
    else:
        print(f"[PASS] 算法重建: {len(alg['entries'])} 入口 / {len(steps)} 步骤，全部带锚点")

    design = load_design_spec(FIXTURE / "algorithms" / "report-export.yaml")
    report = align_algorithms(alg, design)
    st = report["stats"]
    missing_ids = {m.get("design_id") for m in report["missing_in_code"]}
    if "check-permission" not in missing_ids:
        failed = True
        print("[FAIL] 对齐未识别设计缺失步骤 check-permission")
    elif st["matched"] >= 5 and report["review_queue"]:
        print(f"[PASS] 对齐: 匹配 {st['matched']} / 设计缺失 {st['missing_in_code']} / "
              f"审核队列 {st['needs_review']}（含 check-permission）")
    else:
        failed = True
        print(f"[FAIL] 对齐统计异常: {st}")

    print("\n" + ("ALL PASS" if not failed else "HAS FAILURES"))
    return 0 if not failed else 1


def test_report_lineage_and_design_import() -> int:
    import shutil
    spark_fixture = ROOT / "tests" / "fixtures" / "spark-sql-demo"
    for d in ("out", ".report-spliter"):
        shutil.rmtree(spark_fixture / d, ignore_errors=True)

    from rs.lineage import build_registry
    from rs.design_import import import_design

    registry = build_registry(spark_fixture)
    by_name = {r["name"]: r for r in registry["reports"]}
    failed = False

    checks = [
        ("monthly_sales" in by_name, "发现 monthly_sales"),
        ("sales_agg" in by_name, "发现 sales_agg"),
        ("orders" in by_name and by_name["orders"]["kind"] == "csv", "发现 CSV 报告 orders"),
        ("sales_agg" in by_name["monthly_sales"]["upstream_tables"], "monthly_sales 上游含 sales_agg"),
        ({"orders", "customers"} <= set(by_name["sales_agg"]["upstream_tables"]),
         "sales_agg 上游含 orders/customers"),
        ("orders" in by_name["orders"]["upstream_tables"], "CSV orders 上游含 orders 表"),
        (by_name["monthly_sales"]["lineage"].get("monthly_amount") == ["sales_agg.total_amount"],
         "monthly_amount 列级血缘 → sales_agg.total_amount"),
        (registry["dependencies"]["monthly_sales"] == ["sales_agg"], "依赖图: monthly_sales → sales_agg"),
        (registry["dependencies"]["orders"] == [], "CSV orders 无自引用"),
    ]
    for ok, msg in checks:
        if ok:
            print(f"[PASS] {msg}")
        else:
            failed = True
            print(f"[FAIL] {msg}")

    # 设计文档导入：md + 临时 docx/xlsx
    md_steps = import_design(FIXTURE / "algorithms" / "docs" / "report-export-design.md")
    labels = [s["label"] for s in md_steps]
    if len(md_steps) == 6 and "校验导出权限" in labels and \
            md_steps[4]["note"] and "1.15" in md_steps[4]["note"]:
        print("[PASS] Markdown 设计文档导入: 6 步含权限校验与系数备注")
    else:
        failed = True
        print(f"[FAIL] Markdown 导入异常: {md_steps}")

    tmp = Path("/tmp")
    from docx import Document
    from openpyxl import Workbook
    docx_path = tmp / "rs-design-test.docx"
    xlsx_path = tmp / "rs-design-test.xlsx"
    doc = Document()
    doc.add_paragraph("1. 解析日期参数（输入：date，输出：day）", style="List Number")
    doc.add_paragraph("2. 计算报表总额（输入：base，输出：computed）", style="List Number")
    doc.save(docx_path)
    wb = Workbook()
    ws = wb.active
    ws.append(["步骤", "类型", "输入", "输出"])
    ws.append(["解析日期参数", "read", "date", "day"])
    ws.append(["计算报表总额", "transform", "base", "computed"])
    wb.save(xlsx_path)
    d_steps = import_design(docx_path)
    x_steps = import_design(xlsx_path)
    if len(d_steps) == 2 and len(x_steps) == 2:
        print("[PASS] Word/Excel 设计文档导入")
    else:
        failed = True
        print(f"[FAIL] Word/Excel 导入: docx={len(d_steps)} xlsx={len(x_steps)}")
    for p in (docx_path, xlsx_path, Path("/tmp/design-test.docx"), Path("/tmp/design-test.xlsx")):
        try:
            p.unlink()
        except OSError:
            pass

    print("\n" + ("ALL PASS" if not failed else "HAS FAILURES"))
    return 0 if not failed else 1


if __name__ == "__main__":
    rc1 = main()
    rc2 = test_algorithm_and_alignment()
    rc3 = test_report_lineage_and_design_import()
    sys.exit(max(rc1, rc2, rc3))
